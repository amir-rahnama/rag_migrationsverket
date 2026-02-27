"""Extract text from DOCX files using python-docx."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass
class ParsedDocx:
    url: str
    title: str
    text: str
    file_type: str = "docx"


def _clean(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def process_docx(url: str, file_path: str | Path) -> ParsedDocx:
    path = Path(file_path)
    doc = Document(path)

    title = path.stem.replace("_", " ").replace("-", " ")

    # Try to use the first heading as title
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            title = para.text.strip()
            break

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = _clean("\n\n".join(paragraphs))

    return ParsedDocx(url=url, title=title, text=text)
